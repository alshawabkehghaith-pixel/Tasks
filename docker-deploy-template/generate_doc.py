"""
Generates the Docker Deployment Template as a Word document.
Run from inside the docker-deploy-template folder:
    python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ───────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Colour palette ─────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x35, 0x64)   # headings
TEAL   = RGBColor(0x00, 0x70, 0x7F)   # subheadings
GRAY   = RGBColor(0x44, 0x44, 0x44)   # body
CODE_BG = RGBColor(0xF2, 0xF2, 0xF2)  # code block shading
RED    = RGBColor(0xC0, 0x00, 0x00)   # warnings / never-commit notes

# ── Helpers ────────────────────────────────────────────────────

def set_para_shading(para, fill_hex):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    return p

def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = color if color else GRAY
    run.italic = italic
    return p

def note(text):
    """Callout for important notes."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    set_para_shading(p, "EAF4FB")
    run = p.add_run("ℹ  " + text)
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x50, 0x7A)
    return p

def warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    set_para_shading(p, "FFF2CC")
    run = p.add_run("⚠  " + text)
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
    return p

def code_block(lines):
    """Render a list of strings as a shaded monospace block."""
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    set_para_shading(p, "F2F2F2")
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def inline_code(para, text):
    run = para.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return run

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GRAY
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GRAY
    return p

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def table_row(tbl, cells, bold_first=False):
    row = tbl.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY
        if bold_first and i == 0:
            run.bold = True
            run.font.name = "Courier New"


# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("Containerisation & Deployment Guide")
run.bold           = True
run.font.size      = Pt(26)
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Internal Engineering Reference  —  Consulting Agency")
run.font.size      = Pt(12)
run.font.color.rgb = TEAL
run.italic         = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# OVERVIEW
# ══════════════════════════════════════════════════════════════

h1("Overview")
body(
    "This guide walks you through containerising a web application — a Python backend and a "
    "Node.js frontend — behind an Nginx reverse proxy with HTTPS. By the end, the application "
    "will be running on a cloud VM, accessible via your domain with a valid SSL certificate."
)
body(
    "The guide is structured as four sequential phases. Complete them in order. Each phase ends "
    "with a verification step so you know everything is working before moving on."
)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("The four phases are:")
run.font.size = Pt(10.5)
run.font.color.rgb = GRAY

bullet("Phase 1 — Containerise and test the stack locally")
bullet("Phase 2 — Obtain an SSL certificate")
bullet("Phase 3 — Deploy to a cloud VM")
bullet("Phase 4 — Connect your domain and enable HTTPS")

hr()


# ══════════════════════════════════════════════════════════════
# ARCHITECTURE
# ══════════════════════════════════════════════════════════════

h1("How the Stack Works")
body(
    "Before touching any files it helps to understand the architecture, because it explains "
    "every decision made in the template."
)
body(
    "Nginx sits in front of everything. It is the only service that listens on public ports "
    "(80 and 443). When a request arrives, Nginx reads the URL path and decides where to send it:"
)
bullet("Requests to /api/... are forwarded to the backend container")
bullet("Everything else is forwarded to the frontend container")
body(
    "The backend and frontend containers are never exposed to the outside world — only to Nginx, "
    "over Docker's internal network. This means that even if someone tries to hit the backend "
    "port directly from the internet, they cannot."
)
code_block([
    "Internet",
    "  └─► Nginx  :80 / :443          (only public entry point)",
    "        ├─ /api/*  ──► Backend   :{{BACKEND_PORT}}",
    "        └─ /*      ──► Frontend  :{{FRONTEND_PORT}}",
])

hr()


# ══════════════════════════════════════════════════════════════
# TEMPLATE STRUCTURE
# ══════════════════════════════════════════════════════════════

h1("Template Structure")
body("The template folder contains everything you need. Here is what each file does:")
code_block([
    "docker-deploy-template/",
    "├── README.md",
    "├── backend/",
    "│   ├── Dockerfile          Builds the Python backend image",
    "│   ├── .dockerignore       Files excluded from the image",
    "│   ├── .env.example        List of required secrets — commit this",
    "│   └── .env                Actual secrets — NEVER commit this",
    "├── frontend/",
    "│   ├── Dockerfile          Builds the Node.js frontend image",
    "│   └── .dockerignore",
    "├── nginx/",
    "│   ├── nginx.ssl.conf      Nginx config — routes traffic and handles SSL",
    "│   └── certs/              Your SSL certificate files — NEVER commit this",
    "└── docker-compose.yml      Defines all services and how they connect",
])

hr()


# ══════════════════════════════════════════════════════════════
# CUSTOMISATION
# ══════════════════════════════════════════════════════════════

h1("Step 0 — Customise the Template")
body(
    "Every file in the template contains placeholders in the format {{PLACEHOLDER}}. "
    "Replace them all before running anything. The easiest way is a global find-and-replace:"
)
note("In VS Code or Cursor: Ctrl+Shift+H → type the placeholder → type the value → Replace All.")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Placeholders to replace:")
run.font.size = Pt(10.5); run.font.color.rgb = GRAY; run.bold = True

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for cell, text in zip(hdr, ["Placeholder", "What it is", "Example"]):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY

rows = [
    ("{{PYTHON_VERSION}}",       "Python base image tag",                        "3.11"),
    ("{{NODE_VERSION}}",         "Node.js base image tag",                       "20"),
    ("{{BACKEND_PORT}}",         "Port the backend binds to",                    "8081"),
    ("{{FRONTEND_PORT}}",        "Port the frontend binds to",                   "3000"),
    ("{{API_PREFIX}}",           "URL prefix for all API routes",                "/api"),
    ("{{BACKEND_START_CMD}}",    "Command that starts the backend",              "python -m app.server"),
    ("{{BACKEND_INTERNAL_URL}}", "Backend URL seen by the Next.js Node process", "http://backend:8081"),
    ("{{BACKEND_HEALTHCHECK_URL}}", "Full URL for the Compose healthcheck",      "http://localhost:8081/api/healthz"),
    ("{{DOMAIN}}",               "Your public domain name",                      "app.example.com"),
]
for r in rows:
    table_row(tbl, r, bold_first=True)

doc.add_paragraph()  # spacer

hr()


# ══════════════════════════════════════════════════════════════
# REQUIRED APP CHANGES
# ══════════════════════════════════════════════════════════════

h1("Step 0 (continued) — Required Application Code Changes")
body(
    "Three small changes must be made to the application source before building images. "
    "These are not optional — the stack will not work correctly without them."
)

h2("1. Frontend — Make the backend URL configurable")
body(
    "Next.js rewrites typically hardcode localhost as the backend destination. Inside Docker, "
    "services communicate by their service name, not localhost. Update next.config.ts:"
)
code_block([
    "// Before",
    "destination: 'http://localhost:8081/api/:path*'",
    "",
    "// After",
    "const backendUrl = process.env.BACKEND_URL ?? 'http://localhost:8081';",
    "destination: `${backendUrl}/api/:path*`",
])
note(
    "Vite / CRA apps that use browser-side relative paths (/api/...) do not need this change."
)

h2("2. Backend — Update CORS allowed origins")
body(
    "When the app runs behind Nginx on port 80, the browser sends Origin: http://localhost — "
    "not http://localhost:{{FRONTEND_PORT}}. The backend must permit this origin:"
)
code_block([
    "CORS_ALLOWED_ORIGINS = [",
    '    "http://localhost:{{FRONTEND_PORT}}",  # local dev',
    '    "http://localhost",                    # containerised via Nginx',
    '    "http://localhost:80",',
    '    # "https://{{DOMAIN}}",               # add in Phase 4',
    "]",
])

h2("3. Backend — Disable reload mode")
body(
    "Hot-reload (e.g. reload=true in uvicorn) watches the filesystem for changes. Inside a "
    "container the filesystem is static. Leaving it enabled causes instability. "
    "Set reload to false in your backend configuration file before building."
)

hr()


# ══════════════════════════════════════════════════════════════
# PHASE 1
# ══════════════════════════════════════════════════════════════

h1("Phase 1 — Containerise and Test Locally")
body(
    "With the placeholders replaced and the application code updated, you are ready to build "
    "and run the full stack as Docker containers for the first time."
)

h2("Step 1.1 — Create the secrets file")
body(
    "The backend reads credentials from a .env file at startup. Copy the example file and "
    "populate it with real values:"
)
code_block([
    "cp backend/.env.example backend/.env",
    "# Open backend/.env and fill in all values",
])
warning("Never commit backend/.env to version control. It is already listed in .dockerignore.")

h2("Step 1.2 — Build the images")
body(
    "This command reads the Dockerfiles and builds container images for the backend and frontend. "
    "It can take a few minutes the first time."
)
code_block(["docker compose build"])
body("If the build fails, jump to the Debugging section.")

h2("Step 1.3 — Start the stack")
code_block(["docker compose up -d"])
body(
    "The -d flag runs the stack in the background. Docker Compose enforces the startup order "
    "automatically: the backend must pass its healthcheck before the frontend starts, and the "
    "frontend must be running before Nginx starts."
)

h2("Step 1.4 — Verify everything is working")
body("Run these checks one by one:")

numbered("Confirm all three services are running and only Nginx shows published ports:")
code_block(["docker compose ps"])

numbered("Confirm the API responds through Nginx (expect a JSON response):")
code_block(["curl http://localhost{{API_PREFIX}}/healthz"])

numbered("Confirm the backend is not directly accessible (must say connection refused):")
code_block(["curl http://localhost:{{BACKEND_PORT}}/"])

numbered("Confirm the frontend is not directly accessible (must say connection refused):")
code_block(["curl http://localhost:{{FRONTEND_PORT}}/"])

numbered("Open http://localhost in a browser and exercise the application end-to-end.")

numbered("Confirm the stack restarts cleanly:")
code_block(["docker compose down && docker compose up -d && docker compose ps"])

body("Once all six checks pass, Phase 1 is complete.")

h2("Useful commands")
code_block([
    "docker compose logs -f              # stream all logs",
    "docker compose logs <service>       # logs for one service",
    "docker compose restart <service>    # restart after a config change",
    "docker compose down                 # stop and remove containers",
    "docker compose down --rmi all       # full reset including images",
    "docker compose exec <service> sh    # open a shell inside a container",
])

hr()


# ══════════════════════════════════════════════════════════════
# PHASE 2
# ══════════════════════════════════════════════════════════════

h1("Phase 2 — Obtain an SSL Certificate")
body(
    "HTTPS is required for production. You have two options: purchase a certificate from a "
    "Certificate Authority (CA) such as Namecheap, or use Let's Encrypt for free. "
    "Both options are covered below."
)

h2("Option A — Purchased Certificate (recommended)")

h3("Step 2A.1 — Generate a private key and CSR")
body(
    "A Certificate Signing Request (CSR) is a file you submit to the CA to prove you want a "
    "certificate for your domain. It is generated alongside a private key. "
    "Run from the project root:"
)
code_block([
    "# Windows PowerShell",
    'docker run --rm -v "${PWD}/nginx/certs:/certs" alpine/openssl req -new `',
    "  -newkey rsa:2048 -nodes `",
    "  -keyout /certs/your_domain.key `",
    "  -out /certs/your_domain.csr `",
    '  -subj "/CN={{DOMAIN}}/O=YourOrganisation/C=US"',
])
code_block([
    "# Mac / Linux",
    'docker run --rm -v "$(pwd)/nginx/certs:/certs" alpine/openssl req -new \\',
    "  -newkey rsa:2048 -nodes \\",
    "  -keyout /certs/your_domain.key \\",
    "  -out /certs/your_domain.csr \\",
    '  -subj "/CN={{DOMAIN}}/O=YourOrganisation/C=US"',
])
warning("your_domain.key is your private key. It must never leave this machine or be committed to version control.")

h3("Step 2A.2 — Submit the CSR to your CA")
bullet("Log into your CA portal (e.g. Namecheap → SSL Certificates → Activate)")
bullet("Choose Server-Side Automation / Manual CSR")
bullet("Set server software to Nginx")
bullet("Paste the full contents of your_domain.csr")

h3("Step 2A.3 — Complete domain validation")
body(
    "The CA must verify that you control the domain. The CNAME method is recommended "
    "as it does not require the server to be running:"
)
bullet("The CA gives you a CNAME record with a Host value and a Value")
bullet("In your DNS provider, add the CNAME record — enter only the subdomain part in the Host field, the provider appends the domain automatically")
bullet("Wait 5–15 minutes for propagation")
note("Verify propagation at dnschecker.org before clicking Verify in the CA portal.")

h3("Step 2A.4 — Combine the certificate files")
body(
    "The CA emails a zip file containing your_domain.crt and your_domain.ca-bundle. "
    "Extract both into nginx/certs/, then combine them into a single fullchain file:"
)
code_block([
    "# Windows",
    "Get-Content nginx\\certs\\your_domain.crt, nginx\\certs\\your_domain.ca-bundle |",
    "  Set-Content nginx\\certs\\fullchain.crt",
])
code_block([
    "# Mac / Linux",
    "cat nginx/certs/your_domain.crt nginx/certs/your_domain.ca-bundle \\",
    "  > nginx/certs/fullchain.crt",
])
body("nginx/certs/ should now contain two files: your_domain.key and fullchain.crt.")
warning("Add nginx/certs/ to .gitignore before committing anything.")

h2("Option B — Let's Encrypt (Certbot)")
body(
    "Let's Encrypt issues free certificates automatically. The DNS A record for your domain "
    "must already point to a publicly reachable server before you begin."
)

h3("Step 2B.1 — Add Certbot to docker-compose.yml")
code_block([
    "certbot:",
    "  image: certbot/certbot",
    "  volumes:",
    "    - certbot_certs:/etc/letsencrypt",
    "    - certbot_www:/var/www/certbot",
    "",
    "# Also add to the nginx service volumes:",
    "#   - certbot_certs:/etc/letsencrypt:ro",
    "#   - certbot_www:/var/www/certbot:ro",
    "",
    "volumes:",
    "  certbot_certs:",
    "  certbot_www:",
])

h3("Step 2B.2 — Bootstrap and issue the certificate")
body(
    "Nginx cannot start with the SSL config until certificate files exist. "
    "You need to start Nginx on HTTP first, issue the cert, then switch to SSL."
)
body("Create nginx/nginx.certbot-init.conf:")
code_block([
    "events { worker_connections 1024; }",
    "http {",
    "  server {",
    "    listen 80;",
    "    server_name {{DOMAIN}};",
    "    location /.well-known/acme-challenge/ { root /var/www/certbot; }",
    "    location / { return 200 'OK'; add_header Content-Type text/plain; }",
    "  }",
    "}",
])
body("Point the nginx volume to this file temporarily, start Nginx, then issue the certificate:")
code_block([
    "docker compose up -d nginx",
    "",
    "docker compose run --rm certbot certonly \\",
    "  --webroot --webroot-path=/var/www/certbot \\",
    "  --email your@email.com --agree-tos --no-eff-email \\",
    "  -d {{DOMAIN}}",
])
body("Restore nginx.ssl.conf in the volumes and restart Nginx:")
code_block(["docker compose restart nginx"])

h3("Step 2B.3 — Automate renewal")
body("Let's Encrypt certificates expire after 90 days. Add a monthly cron job on the server:")
code_block([
    "# crontab -e",
    "0 3 1 * * docker compose -f /path/to/docker-compose.yml run --rm certbot renew \\",
    "  && docker compose -f /path/to/docker-compose.yml restart nginx",
])

hr()


# ══════════════════════════════════════════════════════════════
# PHASE 3
# ══════════════════════════════════════════════════════════════

h1("Phase 3 — Deploy to a Cloud VM")
body(
    "The application is containerised and tested locally. The certificate is ready. "
    "Now you deploy to a cloud VM so the application is publicly accessible."
)
body(
    "Any Linux VM with a public IP and SSH access works. Ubuntu 22.04 LTS is recommended. "
    "Minimum spec: 1 vCPU, 1 GB RAM. Providers: AWS EC2, GCP Compute Engine, DigitalOcean, "
    "Azure, Oracle Cloud."
)

h2("Step 3.1 — Install Docker on the VM")
body("SSH into the VM, then run:")
code_block([
    "sudo apt update && sudo apt upgrade -y",
    "curl -fsSL https://get.docker.com | sudo sh",
    "sudo usermod -aG docker $USER && newgrp docker",
    "docker --version   # confirm installation",
])

h2("Step 3.2 — Open firewall ports")
body(
    "Traffic will never reach the VM unless two firewall layers are both open. "
    "This is a common source of confusion — both must be configured."
)

h3("Layer 1 — Cloud provider firewall")
body("Add inbound rules allowing TCP port 80 and TCP port 443 from 0.0.0.0/0 (all sources).")

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = "Table Grid"
for cell, text in zip(tbl2.rows[0].cells, ["Provider", "Where to find it"]):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
for row in [
    ("AWS",           "EC2 → Security Groups → Inbound Rules"),
    ("GCP",           "VPC Network → Firewall → Create Firewall Rule"),
    ("DigitalOcean",  "Networking → Firewalls → Inbound Rules"),
    ("Azure",         "Networking → Add Inbound Port Rule"),
    ("Oracle Cloud",  "VCN → Security List → Add Ingress Rules"),
]:
    table_row(tbl2, row)

doc.add_paragraph()

h3("Layer 2 — OS-level firewall (Ubuntu)")
code_block([
    "sudo iptables -I INPUT -p tcp --dport 80 -j ACCEPT",
    "sudo iptables -I INPUT -p tcp --dport 443 -j ACCEPT",
    "sudo apt install -y iptables-persistent",
    "sudo netfilter-persistent save",
])
note(
    "If the site times out but DNS resolves correctly, this OS-level firewall is almost always "
    "the cause. Apply these commands even if the cloud firewall is already open."
)

h2("Step 3.3 — Transfer files to the VM")
body("Choose the option that suits your project size.")

h3("Option A — Transfer the full project (simple projects)")
code_block([
    "# From your local machine",
    "scp -r /path/to/project user@<vm-ip>:~/app",
    "",
    "# On the VM",
    "cd ~/app && docker compose up -d",
])

h3("Option B — Docker Hub + transfer config only (recommended for large images)")
body(
    "Building images on a low-resource VM is slow and may run out of memory. "
    "Build locally, push to Docker Hub, and pull on the VM."
)
body("On your local machine:")
code_block([
    "docker login",
    "docker tag <backend-image>  <hub-username>/<app>-backend:latest",
    "docker tag <frontend-image> <hub-username>/<app>-frontend:latest",
    "docker push <hub-username>/<app>-backend:latest",
    "docker push <hub-username>/<app>-frontend:latest",
])
body("Update docker-compose.yml — replace the build: blocks with image: references:")
code_block([
    "backend:",
    "  image: <hub-username>/<app>-backend:latest",
    "",
    "frontend:",
    "  image: <hub-username>/<app>-frontend:latest",
])
body("Transfer only the config files (not the source code) to the VM:")
code_block([
    "nginx/                  (including the certs/ folder)",
    "docker-compose.yml",
    "backend/.env",
])
body("On the VM:")
code_block(["docker compose pull && docker compose up -d"])

h3("Updating after code changes")
code_block([
    "# Local",
    "docker compose build",
    "docker push <hub-username>/<app>-backend:latest",
    "docker push <hub-username>/<app>-frontend:latest",
    "",
    "# VM",
    "docker compose pull && docker compose up -d --force-recreate",
])

hr()


# ══════════════════════════════════════════════════════════════
# PHASE 4
# ══════════════════════════════════════════════════════════════

h1("Phase 4 — Connect Your Domain and Enable HTTPS")
body(
    "The final step ties everything together: your domain routes to the VM, Nginx serves "
    "the application over HTTPS, and the certificate is trusted by browsers."
)
body("Before starting, confirm:")
bullet("nginx/certs/ contains fullchain.crt and your_domain.key on the VM")
bullet("The DNS A record for {{DOMAIN}} points to the VM's public IP")
bullet("Ports 80 and 443 are open (both firewall layers from Phase 3)")

h2("Step 4.1 — Verify DNS")
code_block(["nslookup {{DOMAIN}}   # must return the VM's public IP"])
note("If DNS is not resolving yet, wait for propagation and check dnschecker.org.")

h2("Step 4.2 — Update nginx.ssl.conf")
body("Replace all {{DOMAIN}} placeholders with your actual domain. Confirm the certificate paths match your filenames:")
code_block([
    "ssl_certificate     /etc/nginx/certs/fullchain.crt;",
    "ssl_certificate_key /etc/nginx/certs/your_domain.key;",
])

h2("Step 4.3 — Update CORS")
body("Remove the localhost entries and add the production domain:")
code_block([
    "CORS_ALLOWED_ORIGINS = [",
    '    "https://{{DOMAIN}}",',
    "]",
])

h2("Step 4.4 — Rebuild and restart")
code_block([
    "docker compose down",
    "docker compose build backend",
    "docker compose up -d",
])

h2("Step 4.5 — Verify HTTPS")
numbered("HTTP should redirect to HTTPS (expect a 301 response):")
code_block(["curl -v http://{{DOMAIN}}/"])

numbered("HTTPS API call should succeed without the -k flag (certificate is trusted):")
code_block(["curl https://{{DOMAIN}}{{API_PREFIX}}/healthz"])

numbered("Backend must still be unreachable directly:")
code_block(["curl https://{{DOMAIN}}:{{BACKEND_PORT}}/   # must refuse connection"])

numbered("Open https://{{DOMAIN}} in a browser — green padlock, no warnings, application loads fully.")

body("All four checks passing means the deployment is complete.")

hr()


# ══════════════════════════════════════════════════════════════
# DEBUGGING
# ══════════════════════════════════════════════════════════════

h1("Debugging")
body("A collection of the most common problems and how to fix them.")

h2("Backend shows 'unhealthy'")
code_block(["docker compose logs backend"])
body(
    "If the application is running but the healthcheck fails, the URL in "
    "{{BACKEND_HEALTHCHECK_URL}} does not match an existing endpoint. "
    "If you see 'curl: command not found', the healthcheck is trying to run curl in a slim "
    "Python image where it doesn't exist — use the Python-based healthcheck already provided "
    "in docker-compose.yml."
)

h2("502 Bad Gateway")
body("Nginx is running but cannot reach an upstream container.")
code_block([
    "docker compose ps",
    "docker compose logs backend",
    "docker compose logs frontend",
])
body("Verify the port numbers in nginx.ssl.conf match what the services actually bind to.")

h2("CORS error in browser console")
body(
    "The browser's Origin header is not listed in the backend's CORS_ALLOWED_ORIGINS. "
    "Apply the change from Step 0 and restart the backend."
)

h2("nginx: [emerg] — SSL error in logs")
code_block(["docker compose logs nginx"])
bullet("fullchain.crt or your_domain.key not present in nginx/certs/ on the server")
bullet("Private key does not match the certificate — the .key file must be from the same run that produced the .csr")
bullet("fullchain.crt is missing the intermediate chain — re-combine your_domain.crt + your_domain.ca-bundle")

h2("Domain times out (DNS resolves correctly)")
body(
    "The VM firewall is blocking inbound traffic. Verify both the cloud provider security "
    "group and the OS-level iptables rules from Phase 3, Step 3.2."
)

h2("pip install — package not found during build")
body(
    "A package name in requirements.txt is incorrect. Common mistakes: "
    "dotenv → python-dotenv, sklearn → scikit-learn, PIL → Pillow."
)

h2("npm ci — missing lockfile")
body(
    "Run npm install locally in the frontend directory to generate package-lock.json, "
    "commit it, then rebuild the image."
)

hr()


# ══════════════════════════════════════════════════════════════
# EXTENSIONS
# ══════════════════════════════════════════════════════════════

h1("Extensions")
body("Drop-in service blocks to add to docker-compose.yml for common stack components.")

h2("PostgreSQL")
code_block([
    "db:",
    "  image: postgres:15",
    "  environment:",
    "    - POSTGRES_USER=${DB_USER}",
    "    - POSTGRES_PASSWORD=${DB_PASSWORD}",
    "    - POSTGRES_DB=${DB_NAME}",
    "  expose:",
    '    - "5432"',
    "  volumes:",
    "    - db_data:/var/lib/postgresql/data",
    "  healthcheck:",
    '    test: ["CMD", "pg_isready", "-U", "${DB_USER}"]',
    "    interval: 10s",
    "    retries: 5",
    "",
    "volumes:",
    "  db_data:",
])
note('Add "depends_on: db: condition: service_healthy" to the backend service.')

h2("Microsoft SQL Server")
code_block([
    "sqlserver:",
    "  image: mcr.microsoft.com/mssql/server:2022-latest",
    "  environment:",
    "    - ACCEPT_EULA=Y",
    "    - SA_PASSWORD=${DB_PASSWORD}",
    "    - MSSQL_PID=Developer",
    "  expose:",
    '    - "1433"',
    "  volumes:",
    "    - sqlserver_data:/var/opt/mssql",
    "  healthcheck:",
    "    test: [\"CMD-SHELL\", \"/opt/mssql-tools/bin/sqlcmd -S localhost -U sa -P ${DB_PASSWORD} -Q 'SELECT 1' || exit 1\"]",
    "    interval: 15s",
    "    retries: 10",
    "    start_period: 30s",
])
note("Uncomment the ODBC driver block in backend/Dockerfile when using pyodbc.")

h2("Apache Airflow (LocalExecutor)")
code_block([
    "airflow-postgres:",
    "  image: postgres:15",
    "  environment:",
    "    - POSTGRES_USER=airflow",
    "    - POSTGRES_PASSWORD=airflow",
    "    - POSTGRES_DB=airflow",
    "  expose:",
    '    - "5432"',
    "  healthcheck:",
    '    test: ["CMD", "pg_isready", "-U", "airflow"]',
    "    interval: 10s",
    "    retries: 5",
    "",
    "airflow-init:",
    "  image: apache/airflow:2.9.0",
    "  depends_on:",
    "    airflow-postgres:",
    "      condition: service_healthy",
    "  environment:",
    "    - AIRFLOW__DATABASE__SQL_ALCHEMY_CONN=postgresql+psycopg2://airflow:airflow@airflow-postgres/airflow",
    "    - AIRFLOW__CORE__FERNET_KEY=${AIRFLOW_FERNET_KEY}",
    "  command: db migrate",
    "",
    "airflow-scheduler:",
    "  image: apache/airflow:2.9.0",
    "  depends_on:",
    "    airflow-init:",
    "      condition: service_completed_successfully",
    "  environment:",
    "    - AIRFLOW__DATABASE__SQL_ALCHEMY_CONN=postgresql+psycopg2://airflow:airflow@airflow-postgres/airflow",
    "    - AIRFLOW__CORE__FERNET_KEY=${AIRFLOW_FERNET_KEY}",
    "  volumes:",
    "    - ./dags:/opt/airflow/dags",
    "    - airflow_logs:/opt/airflow/logs",
    "  command: scheduler",
    "",
    "airflow-webserver:",
    "  image: apache/airflow:2.9.0",
    "  depends_on:",
    "    airflow-init:",
    "      condition: service_completed_successfully",
    "  environment:",
    "    - AIRFLOW__DATABASE__SQL_ALCHEMY_CONN=postgresql+psycopg2://airflow:airflow@airflow-postgres/airflow",
    "    - AIRFLOW__CORE__FERNET_KEY=${AIRFLOW_FERNET_KEY}",
    "    - AIRFLOW__WEBSERVER__BASE_URL=https://{{DOMAIN}}/airflow",
    "  volumes:",
    "    - ./dags:/opt/airflow/dags",
    "    - airflow_logs:/opt/airflow/logs",
    "  expose:",
    '    - "8080"',
    "  command: webserver",
    "  healthcheck:",
    '    test: ["CMD", "curl", "-f", "http://localhost:8080/airflow/health"]',
    "    interval: 15s",
    "    retries: 5",
])
body("Add to nginx.ssl.conf before the catch-all / location:")
code_block([
    "location /airflow/ {",
    "  proxy_pass http://airflow-webserver:8080/airflow/;",
    "  proxy_set_header Host $host;",
    "  proxy_read_timeout 120s;",
    "}",
])
body("Generate the Fernet key once and store it in .env:")
code_block([
    'python -c "from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())"',
])

h2("LDAP (containerised — development only)")
code_block([
    "ldap:",
    "  image: osixia/openldap:1.5.0",
    "  environment:",
    "    - LDAP_ORGANISATION=Example Org",
    "    - LDAP_DOMAIN=example.com",
    "    - LDAP_ADMIN_PASSWORD=${LDAP_ADMIN_PASSWORD}",
    "  expose:",
    '    - "389"',
    '    - "636"',
    "  volumes:",
    "    - ldap_data:/var/lib/ldap",
    "    - ldap_config:/etc/ldap/slapd.d",
])
note("For an external corporate LDAP server, no container is needed — configure the connection via environment variables in backend/.env.")

hr()


# ══════════════════════════════════════════════════════════════
# PRODUCTION CHECKLIST
# ══════════════════════════════════════════════════════════════

h1("Production Checklist")
body("Run through this before going live. Every item must be checked off.")

h2("Security")
bullet("backend/.env is not committed and not present in any image layer")
bullet("nginx/certs/ is in .gitignore")
bullet("Backend and frontend use expose, not ports in docker-compose.yml")
bullet("HTTPS enabled with a valid, trusted certificate")
bullet("HTTP redirects to HTTPS")
bullet("CORS_ALLOWED_ORIGINS contains only the production domain")
bullet("Backend reload / debug mode disabled")
bullet("No placeholder passwords remain in any service")

h2("Reliability")
bullet("All services have a healthcheck defined")
bullet('depends_on uses condition: service_healthy to enforce startup ordering')
bullet("proxy_read_timeout covers the slowest expected operation")
bullet("client_max_body_size covers the largest expected upload")

h2("Maintainability")
bullet("Certificate renewal is automated (Certbot) or calendared (purchased certificate)")
bullet("docker compose down && docker compose up -d works cleanly from any state")
bullet("A new team member can onboard with only: clone repo → populate .env → docker compose up")

hr()


# ══════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ══════════════════════════════════════════════════════════════

h1("Quick Reference")

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Table Grid"
for cell, text in zip(tbl3.rows[0].cells, ["Task", "Command"]):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY

for row in [
    ("Build all images",                  "docker compose build"),
    ("Start stack",                       "docker compose up -d"),
    ("Stop stack",                        "docker compose down"),
    ("View containers",                   "docker compose ps"),
    ("Stream all logs",                   "docker compose logs -f"),
    ("Logs for one service",              "docker compose logs <service>"),
    ("Restart one service",               "docker compose restart <service>"),
    ("Rebuild + restart one service",     "docker compose build <service> && docker compose up -d --force-recreate <service>"),
    ("Shell into container",              "docker compose exec <service> sh"),
    ("Validate Nginx config",             "docker compose exec nginx nginx -t"),
    ("Full reset (removes images)",       "docker compose down --rmi all"),
    ("Generate Fernet key",               'python -c "from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())"'),
]:
    table_row(tbl3, row, bold_first=False)

doc.add_paragraph()

# ── Save ──────────────────────────────────────────────────────
doc.save("Containerisation_and_Deployment_Guide.docx")
print("Done: Containerisation_and_Deployment_Guide.docx")
